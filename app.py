# app.py

```python id="9y1qhp"
import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from openai import OpenAI
import time

# ----------------------------------------
# CONFIGURACIÓN
# ----------------------------------------

st.set_page_config(
    page_title="Resumen Ejecutivo IA",
    page_icon="📄",
    layout="wide"
)

# ----------------------------------------
# ESTILOS
# ----------------------------------------

st.markdown("""
<style>

.main-title {
    font-size:42px;
    font-weight:bold;
    color:#1565C0;
}

.subtitle {
    font-size:18px;
    color:gray;
}

.summary-box {
    background-color:#f5f7fa;
    padding:25px;
    border-radius:15px;
    border:1px solid #dfe6ee;
}

</style>
""", unsafe_allow_html=True)

# ----------------------------------------
# TÍTULO
# ----------------------------------------

st.markdown(
    '<p class="main-title">📄 Generador de Resúmenes Ejecutivos con IA</p>',
    unsafe_allow_html=True
)

st.markdown(
    '<p class="subtitle">Carga uno o varios archivos PDF y obtén un resumen ejecutivo automáticamente.</p>',
    unsafe_allow_html=True
)

# ----------------------------------------
# API KEY
# ----------------------------------------

api_key = st.sidebar.text_input(
    "🔑 OpenAI API Key",
    type="password"
)

# ----------------------------------------
# CARGAR PDFs
# ----------------------------------------

uploaded_files = st.file_uploader(
    "📂 Cargar archivos PDF",
    type=["pdf"],
    accept_multiple_files=True
)

# ----------------------------------------
# EXTRAER TEXTO
# ----------------------------------------

def extraer_texto_pdf(pdf_file):

    texto = ""

    pdf = PdfReader(pdf_file)

    for pagina in pdf.pages:
        contenido = pagina.extract_text()

        if contenido:
            texto += contenido + "\n"

    return texto

# ----------------------------------------
# GENERAR RESUMEN CON IA
# ----------------------------------------

def generar_resumen(texto, api_key):

    client = OpenAI(api_key=api_key)

    prompt = f"""
    Analiza el siguiente documento y genera un resumen ejecutivo profesional.

    Debes incluir:

    1. Objetivo del documento
    2. Puntos principales
    3. Conclusiones
    4. Recomendaciones

    Documento:
    {texto[:12000]}
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": "Eres un analista ejecutivo profesional."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.3
    )

    return response.choices[0].message.content

# ----------------------------------------
# EXPORTAR WORD
# ----------------------------------------

def generar_word(resumen):

    doc = Document()

    doc.add_heading('Resumen Ejecutivo', level=1)

    doc.add_paragraph(resumen)

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer

# ----------------------------------------
# BOTÓN GENERAR
# ----------------------------------------

if st.button("🚀 Generar Resumen Ejecutivo"):

    if not api_key:
        st.error("Ingrese su OpenAI API Key.")
        st.stop()

    if not uploaded_files:
        st.error("Debe cargar al menos un PDF.")
        st.stop()

    progreso = st.progress(0)

    texto_total = ""

    for i, archivo in enumerate(uploaded_files):

        st.info(f"Procesando: {archivo.name}")

        texto_total += extraer_texto_pdf(archivo)

        progreso.progress((i + 1) / len(uploaded_files))

        time.sleep(0.5)

    with st.spinner("🤖 Generando resumen con IA..."):

        resumen = generar_resumen(texto_total, api_key)

    st.success("✅ Resumen generado correctamente")

    # Mostrar resumen
    st.markdown('<div class="summary-box">', unsafe_allow_html=True)

    st.markdown(resumen)

    st.markdown('</div>', unsafe_allow_html=True)

    # Descargar Word
    archivo_word = generar_word(resumen)

    st.download_button(
        label="📥 Descargar Resumen en Word",
        data=archivo_word,
        file_name="resumen_ejecutivo.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Animación
    st.balloons()
```

# requirements.txt

```txt id="2h9b3k"
streamlit
openai
PyPDF2
python-docx
```

# README.md

````md id="dwyxms"
# Generador de Resúmenes Ejecutivos con IA

Aplicación desarrollada con Streamlit para analizar documentos PDF y generar resúmenes ejecutivos automáticos utilizando inteligencia artificial.

## Funcionalidades

- Carga múltiples PDFs
- Extrae texto automáticamente
- Genera resúmenes ejecutivos con IA
- Exporta a Word
- Diseño moderno
- Compatible con Streamlit Cloud

---

## Instalación local

```bash
pip install -r requirements.txt
streamlit run app.py
````

---

## Configuración OpenAI

Necesitas una API Key de OpenAI.

Puedes obtenerla en:

https://platform.openai.com/api-keys

---

## Despliegue en Streamlit Cloud

1. Subir el proyecto a GitHub
2. Entrar a https://share.streamlit.io
3. Conectar GitHub
4. Seleccionar el repositorio
5. Elegir:

   * Branch: main
   * app.py
6. Deploy 🚀

````

# ESTRUCTURA DEL PROYECTO

```bash id="7t5r1x"
resumen-ejecutivo-ia/
│
├── app.py
├── requirements.txt
└── README.md
````

# COMANDOS GITHUB

```bash id="juxd8p"
git init
git add .
git commit -m "Primer commit"
git branch -M main
git remote add origin https://github.com/TU_USUARIO/TU_REPOSITORIO.git
git push -u origin main
```
